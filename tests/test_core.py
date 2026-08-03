"""Regression tests for deterministic interview-coach behavior."""

from __future__ import annotations

import shutil
import tempfile
import unittest
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

from docx import Document

from utils.communication import analyze_communication
from utils import database
from utils.pdf_reader import ResumeExtractionError, extract_text
from utils.question_generator import parse_interview_questions
from utils.report import build_report
from utils.speech import _configure_ffmpeg
from utils.summarizer import summarize_resume
from utils.question_generator import generate_interview_questions


def _record(session_id, question_index, score):
    return {
        "session_id": session_id,
        "question_index": question_index,
        "question": f"Question {question_index}",
        "answer": "A concise answer",
        "answer_source": "Typed",
        "communication": {
            "communication_score": 8,
            "total_fillers": 1,
        },
        "feedback": {
            "score": score,
            "strengths": ["Specific"],
            "improvements": ["Add detail"],
            "better_structure": ["Context", "Action", "Result"],
        },
    }


class CommunicationTests(unittest.TestCase):
    def test_filler_words_and_speaking_rate_are_measured(self):
        result = analyze_communication(
            "Um, I actually built the project and tested it.",
            duration_seconds=6,
        )

        self.assertEqual(result["total_fillers"], 2)
        self.assertEqual(result["word_count"], 9)
        self.assertEqual(result["words_per_minute"], 90)


class QuestionParsingTests(unittest.TestCase):
    def test_numbered_questions_are_parsed(self):
        response = "\n".join(
            f"{index}. Question number {index}?"
            for index in range(1, 6)
        )

        questions = parse_interview_questions(response)

        self.assertEqual(len(questions), 5)
        self.assertEqual(questions[0], "Question number 1?")

    def test_wrong_question_count_is_rejected(self):
        with self.assertRaises(ValueError):
            parse_interview_questions("1. Only one question?")

    @patch("utils.question_generator.generate_text")
    def test_job_description_is_used_for_questions(self, generate):
        generate.return_value = "\n".join(
            f"{index}. Question number {index}?"
            for index in range(1, 6)
        )

        generate_interview_questions(
            "Candidate knows Python.",
            "Data Analyst",
            "Entry level",
            "Requires SQL, dashboards, and stakeholder communication.",
        )

        prompt = generate.call_args.args[0]
        self.assertIn("Requires SQL, dashboards", prompt)
        self.assertIn("BEGIN JOB DESCRIPTION", prompt)


class ResumeSummaryPromptTests(unittest.TestCase):
    @patch("utils.summarizer.generate_text", return_value="Role briefing")
    def test_job_description_is_compared_with_resume(self, generate):
        result = summarize_resume(
            "Built a Python reporting dashboard.",
            "Data Analyst",
            "Requires SQL and stakeholder presentations.",
        )

        self.assertEqual(result, "Role briefing")
        prompt = generate.call_args.args[0]
        self.assertIn("Requires SQL and stakeholder presentations.", prompt)
        self.assertIn("confirmed matches", prompt)


class _UploadedFile:
    def __init__(self, data, name):
        self._data = data
        self.name = name

    def getvalue(self):
        return self._data


class ResumeExtractionTests(unittest.TestCase):
    def test_txt_resume_is_cleaned(self):
        resume = _UploadedFile(
            (
                b"Alex Example\nPython   developer with internship experience. "
                b"Built dashboards, automated reports, and tested a student project."
            ),
            "resume.txt",
        )
        self.assertIn("Python developer", extract_text(resume))

    def test_docx_paragraphs_and_tables_are_extracted(self):
        buffer = BytesIO()
        document = Document()
        document.add_paragraph("Alex Example - Software Engineering Student")
        document.add_paragraph(
            "Built and tested a Python API for a capstone analytics project."
        )
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Skills: Python, SQL, Git, communication"
        document.save(buffer)

        text = extract_text(_UploadedFile(buffer.getvalue(), "resume.docx"))

        self.assertIn("Software Engineering Student", text)
        self.assertIn("Skills: Python, SQL", text)

    def test_too_little_text_has_an_actionable_error(self):
        with self.assertRaisesRegex(ResumeExtractionError, "Very little text"):
            extract_text(_UploadedFile(b"Name only", "resume.txt"))

    @patch("utils.pdf_reader.extract_image_text")
    def test_png_resume_uses_gemini_image_transcription(self, extract_image):
        extract_image.return_value = (
            "Alex Example\nSoftware Engineering Student with Python experience. "
            "Built and tested an interview coaching application for a hackathon."
        )

        text = extract_text(_UploadedFile(b"valid-image-bytes", "resume.png"))

        extract_image.assert_called_once_with(b"valid-image-bytes", "image/png")
        self.assertIn("Software Engineering Student", text)

    @patch("utils.pdf_reader.extract_image_text")
    def test_jpeg_resume_uses_the_jpeg_mime_type(self, extract_image):
        extract_image.return_value = (
            "Alex Example\nData analyst with SQL and dashboard experience. "
            "Automated reporting and presented findings to a student team."
        )

        extract_text(_UploadedFile(b"valid-image-bytes", "resume.jpeg"))

        extract_image.assert_called_once_with(b"valid-image-bytes", "image/jpeg")

    @patch("utils.pdf_reader.extract_image_text")
    def test_image_missing_api_key_has_a_specific_error(self, extract_image):
        extract_image.side_effect = ValueError(
            "GEMINI_API_KEY is missing from .env or Streamlit secrets."
        )

        with self.assertRaisesRegex(ResumeExtractionError, "Gemini is not configured"):
            extract_text(_UploadedFile(b"valid-image-bytes", "resume.png"))


class ReportTests(unittest.TestCase):
    def test_report_aggregates_scores(self):
        report = build_report(
            {
                0: _record("session-a", 0, 9),
                1: _record("session-a", 1, 5),
            }
        )

        self.assertEqual(report["answers_evaluated"], 2)
        self.assertEqual(report["average_score"], 7.0)
        self.assertEqual(report["average_communication"], 8.0)
        self.assertEqual(report["total_fillers"], 2)


class DatabasePrivacyTests(unittest.TestCase):
    def test_history_is_scoped_to_one_session(self):
        with tempfile.TemporaryDirectory() as temporary_directory:
            database_path = Path(temporary_directory) / "interviews.db"
            with patch.object(database, "DATABASE_PATH", database_path):
                database.initialize_database()
                database.save_answer(_record("session-a", 0, 8))
                database.save_answer(_record("session-b", 0, 4))

                history = database.get_history("session-a")

        self.assertEqual(len(history), 1)
        self.assertEqual(history[0]["session_id"], "session-a")
        self.assertEqual(history[0]["score"], 8)


class SpeechDependencyTests(unittest.TestCase):
    def test_ffmpeg_is_exposed_under_whispers_command_name(self):
        executable = _configure_ffmpeg()

        self.assertTrue(executable.exists())
        self.assertIsNotNone(shutil.which("ffmpeg"))


if __name__ == "__main__":
    unittest.main()
