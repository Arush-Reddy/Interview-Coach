from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader
from utils.gemini_client import extract_image_text


MAX_FILE_BYTES = 8 * 1024 * 1024
MAX_TEXT_CHARACTERS = 30_000
MIN_TEXT_CHARACTERS = 80


class ResumeExtractionError(ValueError):
    """An actionable error that is safe to display in the Streamlit UI."""


def _clean_text(text):
    lines = (" ".join(line.split()) for line in text.replace("\x00", "").splitlines())
    return "\n".join(line for line in lines if line).strip()


def _extract_pdf(file_bytes):
    try:
        reader = PdfReader(BytesIO(file_bytes))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as error:
        raise ResumeExtractionError(
            "This PDF could not be read. Export it again or upload a DOCX/TXT copy."
        ) from error


def _extract_docx(file_bytes):
    try:
        document = Document(BytesIO(file_bytes))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        table_cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        return "\n".join(paragraphs + table_cells)
    except Exception as error:
        raise ResumeExtractionError(
            "This DOCX could not be read. Check that it is a valid Word document."
        ) from error


def extract_text(document_file, document_kind="resume"):
    """Extract normalized text from an uploaded document or image."""
    file_bytes = document_file.getvalue()
    filename = getattr(document_file, "name", f"{document_kind}.pdf")
    label = document_kind.strip() or "document"

    if not file_bytes:
        raise ResumeExtractionError(f"The uploaded {label} is empty.")
    if len(file_bytes) > MAX_FILE_BYTES:
        raise ResumeExtractionError(
            f"The {label} is too large. Upload a file under 8 MB."
        )

    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        raw_text = _extract_pdf(file_bytes)
    elif suffix == ".docx":
        raw_text = _extract_docx(file_bytes)
    elif suffix == ".txt":
        try:
            raw_text = file_bytes.decode("utf-8-sig")
        except UnicodeDecodeError as error:
            raise ResumeExtractionError("TXT resumes must use UTF-8 encoding.") from error
    elif suffix in {".png", ".jpg", ".jpeg"}:
        mime_type = "image/png" if suffix == ".png" else "image/jpeg"
        try:
            raw_text = extract_image_text(file_bytes, mime_type)
        except ValueError as error:
            if "GEMINI_API_KEY" in str(error):
                raise ResumeExtractionError(
                    f"Gemini is not configured to read {label} images."
                ) from error
            raise ResumeExtractionError(
                f"Gemini could not read this {label} image. Upload a clear, upright "
                "PNG or JPG."
            ) from error
        except Exception as error:
            raise ResumeExtractionError(
                f"Gemini could not read this {label} image. Check your API access and "
                "upload a clear, upright PNG or JPG."
            ) from error
    else:
        raise ResumeExtractionError("Upload a PDF, DOCX, TXT, PNG, or JPG resume.")

    text = _clean_text(raw_text)
    if len(text) < MIN_TEXT_CHARACTERS:
        raise ResumeExtractionError(
            f"Very little text was found in the {label}. Scanned PDFs need OCR; "
            "try DOCX, TXT, PNG, or JPG."
        )
    return text[:MAX_TEXT_CHARACTERS]
